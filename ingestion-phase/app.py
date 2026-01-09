import streamlit as st
import pandas as pd
import torch
import chromadb
import requests
import json
import os
from pymongo import MongoClient
from bson import ObjectId
from typing import Dict, List, Optional
from dotenv import load_dotenv
import math

# Load environment variables from .env file
load_dotenv()

# Try to import transformers, but make it optional
try:
    from transformers import AutoTokenizer, AutoModel
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    TRANSFORMERS_AVAILABLE = False
    AutoTokenizer = None
    AutoModel = None
import time
from datetime import datetime
import os
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from io import BytesIO
import hashlib
from functools import lru_cache
import httpx

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.units import inch
from pypdf import PdfReader, PdfWriter
from docx import Document

def extract_template_outline(template_bytes: bytes) -> List[str]:
    """Module-level extractor for PDF template headings to avoid class reload ordering issues."""
    try:
        reader = PdfReader(BytesIO(template_bytes))
        text = []
        for page in reader.pages[:3]:
            try:
                text.append(page.extract_text() or "")
            except Exception:
                continue
        joined = "\n".join(text)
        lines = [l.strip() for l in joined.splitlines()]
        candidates: List[str] = []
        for line in lines:
            if not line:
                continue
            if len(line) < 3 or len(line) > 80:
                continue
            if line.lower().startswith("page "):
                continue
            looks_like_heading = (
                line.endswith(":") or
                (line.isupper() and any(c.isalpha() for c in line)) or
                (line.istitle() and sum(ch.isalpha() for ch in line) >= 6)
            )
            if looks_like_heading:
                normalized = line.rstrip(":").strip()
                if normalized not in candidates:
                    candidates.append(normalized)
        return candidates[:30] if candidates else []
    except Exception:
        return []

# Try to import autogen, but make it optional
try:
    import pyautogen
    AUTOGEN_AVAILABLE = True
except ImportError:
    AUTOGEN_AVAILABLE = False
    st.warning("‚ö†Ô∏è AutoGen not available. Some features may be limited.")

# Performance: HTTP session and model/db connectors
def _get_css(minimal: bool) -> str:
    if minimal:
        return """
<style>
    :root {
        --bg:#ffffff; 
        --fg:#0f172a; 
        --muted:#64748b; 
        --border:#e2e8f0; 
        --primary:#0ea5e9; 
        --primary-dark:#0284c7;
        --accent-blue:#0ea5e9;
        --text-dark:#0f172a;
        --text-light:#64748b;
        --success-green:#16a34a;
        --danger-red:#dc2626;
        --card:#ffffff;
        --card-muted:#f8fafc;
    }
    .main-header { padding: 1rem; border: 1px solid var(--border); border-radius: 12px; background: var(--card); color: var(--fg); }
    .main-header h1 { margin:0; font-size: 1.4rem; }
    .patient-card, .metric-card, .chat-container, .summary-card { border: 1px solid var(--border); border-radius: 12px; padding: 1rem; background: var(--card); color: var(--fg); }
    .patient-card h4, .metric-card h4, .summary-card h3 { color: var(--text-dark); }
    .chat-message { border:1px solid var(--border); border-left:4px solid var(--primary); border-radius:10px; padding:.75rem; background:var(--card-muted); color: var(--text-dark); }
    .doctor-message { background:var(--card-muted); }
    .ai-message { background:var(--card-muted); border-left-color:#9333ea; }
    .stButton > button { background: var(--primary); color:#fff; border:0; border-radius:10px; padding:.6rem 1rem; box-shadow: 0 1px 2px rgba(0,0,0,.05); }
    .stButton > button:hover { background: var(--primary-dark); }
    .stTextArea textarea, .stTextInput input { border-radius:10px !important; border:1px solid var(--border) !important; }
</style>
"""
    return ""

from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from requests import Session

def _http_session() -> Session:
    session = requests.Session()
    retries = Retry(total=2, backoff_factor=0.2, status_forcelist=[429, 502, 503, 504])
    adapter = HTTPAdapter(max_retries=retries, pool_connections=10, pool_maxsize=10)
    session.headers.update({"Connection": "keep-alive"})
    session.mount("http://", adapter)
    session.mount("https://", adapter)
    return session

# Embedding cache to avoid recomputing embeddings
_embedding_cache = {}

def _get_text_hash(text: str) -> str:
    """Generate hash for text to use as cache key"""
    return hashlib.md5(text.encode()).hexdigest()

def clean_patient_data_for_json(patient_data: Optional[Dict]) -> Optional[Dict]:
    """Clean patient data for JSON serialization - handle ObjectId, NaN, etc."""
    if patient_data is None:
        return None
    
    cleaned = {}
    for key, value in patient_data.items():
        # Handle ObjectId
        if isinstance(value, ObjectId):
            cleaned[key] = str(value)
        # Handle NaN values
        elif isinstance(value, float) and (math.isnan(value) or math.isinf(value)):
            cleaned[key] = None
        # Handle nested dictionaries
        elif isinstance(value, dict):
            cleaned[key] = clean_patient_data_for_json(value)
        # Handle lists
        elif isinstance(value, list):
            cleaned[key] = [
                clean_patient_data_for_json(item) if isinstance(item, dict) else 
                (str(item) if isinstance(item, ObjectId) else 
                 (None if isinstance(item, float) and (math.isnan(item) or math.isinf(item)) else item))
                for item in value
            ]
        # Handle other types
        else:
            cleaned[key] = value
    return cleaned

# FastAPI client configuration
# FastAPI URL - use environment variable or detect based on deployment
# In production (Docker), FastAPI runs on localhost:8000 internally
# In local dev, use localhost:8000
FASTAPI_BASE_URL = os.getenv("FASTAPI_URL", os.getenv("FASTAPI_BASE_URL", "http://localhost:8000"))

class FastAPIClient:
    """Client for FastAPI backend"""
    def __init__(self, base_url: str = FASTAPI_BASE_URL):
        self.base_url = base_url
        self.client = httpx.Client(timeout=None)  # No timeout - allow long-running LLM requests
    
    def chat(self, message: str, patient_data: Optional[Dict] = None) -> str:
        """Chat with AI agent via FastAPI"""
        try:
            # Clean patient_data to handle ObjectId and other non-serializable types
            cleaned_patient_data = clean_patient_data_for_json(patient_data)
            response = self.client.post(
                f"{self.base_url}/api/chat",
                json={"message": message, "patient_data": cleaned_patient_data}
            )
            response.raise_for_status()
            return response.json()["response"]
        except httpx.RequestError as e:
            return f"‚ùå Error connecting to API: {str(e)}"
        except httpx.HTTPStatusError as e:
            return f"‚ùå API error: {e.response.text}"
    
    def generate_summary(self, patient_data: str, template_outline: Optional[List[str]] = None) -> str:
        """Generate discharge summary via FastAPI"""
        try:
            response = self.client.post(
                f"{self.base_url}/api/generate-summary",
                json={"patient_data": patient_data, "template_outline": template_outline}
            )
            response.raise_for_status()
            return response.json()["summary"]
        except httpx.RequestError as e:
            return f"‚ùå Error connecting to API: {str(e)}"
        except httpx.HTTPStatusError as e:
            return f"‚ùå API error: {e.response.text}"
    
    def search_similar(self, query_text: str, n_results: int = 3) -> List[Dict]:
        """Search similar cases via FastAPI"""
        try:
            response = self.client.post(
                f"{self.base_url}/api/search-similar",
                json={"query_text": query_text, "n_results": n_results}
            )
            response.raise_for_status()
            return response.json()["similar_cases"]
        except httpx.RequestError as e:
            st.error(f"‚ùå Error connecting to API: {str(e)}")
            return []
        except httpx.HTTPStatusError as e:
            st.error(f"‚ùå API error: {e.response.text}")
            return []
    
    def get_patient(self, unit_no: str) -> Optional[Dict]:
        """Get patient via FastAPI"""
        try:
            response = self.client.post(
                f"{self.base_url}/api/patient",
                json={"unit_no": unit_no}
            )
            response.raise_for_status()
            return response.json()["patient"]
        except httpx.RequestError as e:
            st.error(f"‚ùå Error connecting to API: {str(e)}")
            return None
        except httpx.HTTPStatusError as e:
            if e.response.status_code == 404:
                return None
            st.error(f"‚ùå API error: {e.response.text}")
            return None
    
    def get_all_patients(self) -> List[Dict]:
        """Get list of all patients via FastAPI"""
        try:
            response = self.client.get(f"{self.base_url}/api/patients")
            response.raise_for_status()
            return response.json().get("patients", [])
        except httpx.RequestError as e:
            st.error(f"‚ùå Error connecting to API: {str(e)}")
            return []
        except httpx.HTTPStatusError as e:
            st.error(f"‚ùå API error: {e.response.text}")
            return []
    
    def health_check(self) -> bool:
        """Check if FastAPI backend is available"""
        try:
            # Use a longer timeout for first check (model loading can be slow)
            response = self.client.get(f"{self.base_url}/health", timeout=10.0)
            return response.status_code == 200
        except Exception as e:
            # Log the error for debugging but don't fail
            return False
    
    def close(self):
        """Close the HTTP client"""
        self.client.close()

def _load_tokenizer_model():
    if not TRANSFORMERS_AVAILABLE:
        st.error("‚ùå Transformers library not available. Please install it: pip install transformers")
        return None, None
    try:
        tokenizer = AutoTokenizer.from_pretrained("emilyalsentzer/Bio_ClinicalBERT")
        model = AutoModel.from_pretrained("emilyalsentzer/Bio_ClinicalBERT")
        model.eval()
        if torch.cuda.is_available():
            model.to("cuda")
        return tokenizer, model
    except Exception as e:
        st.error(f"‚ùå Error loading transformers model: {str(e)}")
        return None, None

def _connect_mongo(uri: str):
    client = MongoClient(uri)
    return client

def _connect_chroma(path: str):
    client = chromadb.PersistentClient(path=path)
    collection = client.get_or_create_collection("patient_embeddings")
    return client, collection

# Page configuration
st.set_page_config(
    page_title="Medical Discharge Summary Assistant",
    page_icon="üè•",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&family=Poppins:wght@400;500;600;700;800&family=Space+Grotesk:wght@400;500;600;700&display=swap');
    
    /* Ultra Modern Color Palette with Glassmorphism */
    :root {
        --bg-primary: #0a0e27;
        --bg-secondary: #141b2d;
        --bg-tertiary: #1a2332;
        --bg-card: rgba(30, 41, 59, 0.7);
        --bg-card-hover: rgba(36, 52, 71, 0.9);
        --bg-input: rgba(15, 23, 42, 0.8);
        --text-primary: #f1f5f9;
        --text-secondary: #cbd5e1;
        --text-muted: #94a3b8;
        --border-color: rgba(51, 65, 85, 0.5);
        --border-light: rgba(71, 85, 105, 0.6);
        --primary: #6366f1;
        --primary-hover: #4f46e5;
        --primary-light: #818cf8;
        --accent-purple: #a855f7;
        --accent-cyan: #06b6d4;
        --success: #10b981;
        --warning: #f59e0b;
        --danger: #ef4444;
        --gradient-primary: linear-gradient(135deg, #6366f1 0%, #8b5cf6 50%, #ec4899 100%);
        --gradient-secondary: linear-gradient(135deg, #0ea5e9 0%, #3b82f6 100%);
        --gradient-success: linear-gradient(135deg, #10b981 0%, #059669 100%);
        --gradient-card: linear-gradient(135deg, rgba(30, 41, 59, 0.8) 0%, rgba(26, 35, 50, 0.9) 100%);
        --gradient-glass: linear-gradient(135deg, rgba(255, 255, 255, 0.1) 0%, rgba(255, 255, 255, 0.05) 100%);
        --shadow-sm: 0 2px 8px rgba(0, 0, 0, 0.4);
        --shadow-md: 0 8px 24px rgba(0, 0, 0, 0.5);
        --shadow-lg: 0 16px 48px rgba(0, 0, 0, 0.6);
        --shadow-glow: 0 0 30px rgba(99, 102, 241, 0.4);
        --shadow-glow-purple: 0 0 30px rgba(168, 85, 247, 0.4);
        --shadow-glow-cyan: 0 0 30px rgba(6, 182, 212, 0.4);
    }
    
    /* Global Styles with Smooth Animations */
    * {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
    }
    
    /* Hide scrollbar on main page - only chat container shows scrollbar */
    html, body {
        overflow-x: hidden;
        overflow-y: auto;
        height: 100%;
        /* Hide scrollbar but keep scrolling functionality */
        scrollbar-width: none; /* Firefox */
        -ms-overflow-style: none; /* IE and Edge */
    }
    
    html::-webkit-scrollbar, body::-webkit-scrollbar {
        display: none; /* Chrome, Safari, Opera */
    }
    
    .stApp {
        background: linear-gradient(135deg, #0a0e27 0%, #141b2d 50%, #0a0e27 100%);
        background-size: 200% 200%;
        animation: gradientShift 15s ease infinite;
        color: var(--text-primary);
        position: relative;
        overflow-y: auto;
        overflow-x: hidden;
        /* Hide scrollbar but keep scrolling functionality */
        scrollbar-width: none; /* Firefox */
        -ms-overflow-style: none; /* IE and Edge */
    }
    
    .stApp::-webkit-scrollbar {
        display: none; /* Chrome, Safari, Opera */
    }
    
    .stApp::before {
        content: '';
        position: fixed;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: 
            radial-gradient(circle at 20% 30%, rgba(99, 102, 241, 0.1) 0%, transparent 50%),
            radial-gradient(circle at 80% 70%, rgba(168, 85, 247, 0.1) 0%, transparent 50%),
            radial-gradient(circle at 50% 50%, rgba(6, 182, 212, 0.05) 0%, transparent 50%);
        pointer-events: none;
        z-index: 0;
    }
    
    @keyframes gradientShift {
        0%, 100% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
    }
    
    /* Ensure all content is above background - hide scrollbar on main containers */
    .main {
        position: relative;
        z-index: 10 !important;
        overflow-y: auto;
        overflow-x: hidden !important;
        /* Hide scrollbar but keep scrolling functionality */
        scrollbar-width: none; /* Firefox */
        -ms-overflow-style: none; /* IE and Edge */
    }
    
    .main::-webkit-scrollbar {
        display: none; /* Chrome, Safari, Opera */
    }
    
    .main .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
        position: relative;
        z-index: 10 !important;
        overflow-y: visible !important;
        overflow-x: hidden !important;
    }
    
    /* Hide scrollbar on Streamlit app view container */
    [data-testid="stAppViewContainer"] {
        overflow-y: auto;
        overflow-x: hidden !important;
        /* Hide scrollbar but keep scrolling functionality */
        scrollbar-width: none; /* Firefox */
        -ms-overflow-style: none; /* IE and Edge */
    }
    
    [data-testid="stAppViewContainer"]::-webkit-scrollbar {
        display: none; /* Chrome, Safari, Opera */
    }
    
    /* Ensure Streamlit elements are visible */
    [data-testid="stAppViewContainer"],
    [data-testid="stAppViewContainer"] > div,
    .element-container,
    .stMarkdown,
    div[data-testid="stMarkdownContainer"] {
        position: relative;
        z-index: 10 !important;
    }
    
    /* Ensure header is visible */
    .main-header {
        position: relative;
        z-index: 10 !important;
    }
    
    /* Ultra Modern Header with Animated Gradient */
    .main-header {
        padding: 4rem 3rem;
        border: none;
        border-radius: 32px;
        background: var(--gradient-primary);
        color: white;
        text-align: center;
        margin-bottom: 3rem;
        box-shadow: var(--shadow-lg), var(--shadow-glow), 0 0 60px rgba(99, 102, 241, 0.3);
        position: relative;
        overflow: hidden;
        backdrop-filter: blur(20px);
        border: 1px solid rgba(255, 255, 255, 0.1);
        animation: headerFloat 6s ease-in-out infinite;
    }
    
    @keyframes headerFloat {
        0%, 100% { transform: translateY(0px); }
        50% { transform: translateY(-5px); }
    }
    
    .main-header::before {
        content: '';
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: radial-gradient(circle, rgba(255, 255, 255, 0.15) 0%, transparent 70%);
        animation: rotate 20s linear infinite;
        pointer-events: none;
    }
    
    @keyframes rotate {
        from { transform: rotate(0deg); }
        to { transform: rotate(360deg); }
    }
    
    .main-header::after {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: linear-gradient(135deg, rgba(255, 255, 255, 0.1) 0%, transparent 100%);
        pointer-events: none;
    }
    
    .main-header h1 {
        margin: 0;
        font-size: 3rem;
        font-weight: 900;
        letter-spacing: -0.04em;
        position: relative;
        z-index: 1;
        text-shadow: 0 4px 20px rgba(0, 0, 0, 0.4), 0 0 40px rgba(255, 255, 255, 0.2);
        background: linear-gradient(135deg, #ffffff 0%, rgba(255, 255, 255, 0.9) 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        animation: textShine 3s ease-in-out infinite;
    }
    
    @keyframes textShine {
        0%, 100% { filter: brightness(1); }
        50% { filter: brightness(1.2); }
    }
    
    .main-header p {
        margin: 1rem 0 0 0;
        opacity: 0.95;
        font-size: 1.25rem;
        font-weight: 500;
        position: relative;
        z-index: 1;
        text-shadow: 0 2px 10px rgba(0, 0, 0, 0.3);
        letter-spacing: 0.02em;
    }
    
    /* Ultra Modern Glassmorphism Cards */
    .patient-card, .metric-card, .chat-container, .summary-card {
        border: 1px solid rgba(255, 255, 255, 0.1);
        border-radius: 24px;
        padding: 2.5rem;
        background: var(--gradient-card);
        backdrop-filter: blur(20px) saturate(180%);
        -webkit-backdrop-filter: blur(20px) saturate(180%);
        color: var(--text-primary);
        box-shadow: var(--shadow-md), inset 0 1px 0 rgba(255, 255, 255, 0.1);
        transition: all 0.5s cubic-bezier(0.4, 0, 0.2, 1);
        position: relative;
        overflow: hidden;
    }
    
    .patient-card::before, .summary-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(255, 255, 255, 0.1), transparent);
        transition: left 0.5s;
    }
    
    .patient-card:hover::before, .summary-card:hover::before {
        left: 100%;
    }
    
    .patient-card:hover, .metric-card:hover, .summary-card:hover {
        box-shadow: var(--shadow-lg), var(--shadow-glow), 0 0 40px rgba(99, 102, 241, 0.3);
        transform: translateY(-8px) scale(1.02);
        border-color: rgba(99, 102, 241, 0.5);
    }
    
    .metric-card {
        text-align: center;
        background: var(--gradient-card);
        backdrop-filter: blur(20px) saturate(180%);
        border: 1px solid rgba(255, 255, 255, 0.1);
        position: relative;
        overflow: hidden;
        cursor: pointer;
    }
    
    .metric-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 4px;
        background: var(--gradient-primary);
        box-shadow: 0 0 20px rgba(99, 102, 241, 0.6);
    }
    
    .metric-card::after {
        content: '';
        position: absolute;
        top: 50%;
        left: 50%;
        width: 0;
        height: 0;
        border-radius: 50%;
        background: radial-gradient(circle, rgba(99, 102, 241, 0.2) 0%, transparent 70%);
        transform: translate(-50%, -50%);
        transition: width 0.6s, height 0.6s;
    }
    
    .metric-card:hover::after {
        width: 300px;
        height: 300px;
    }
    
    .metric-card h4 {
        margin: 0 0 1rem 0;
        color: var(--primary-light);
        font-size: 1.5rem;
        font-weight: 800;
        letter-spacing: -0.02em;
        position: relative;
        z-index: 1;
        text-shadow: 0 2px 10px rgba(99, 102, 241, 0.3);
    }
    
    .metric-card p {
        margin: 0.5rem 0;
        color: var(--text-secondary);
        font-weight: 600;
        font-size: 1rem;
        position: relative;
        z-index: 1;
    }
    
    /* Ultra Modern Chat Messages with Glassmorphism */
    .chat-message {
        border: 1px solid rgba(255, 255, 255, 0.1);
        border-left: 5px solid var(--primary);
        border-radius: 20px;
        padding: 1.5rem;
        background: rgba(30, 41, 59, 0.6);
        backdrop-filter: blur(15px) saturate(180%);
        color: var(--text-primary);
        margin: 1.25rem 0;
        box-shadow: var(--shadow-sm), inset 0 1px 0 rgba(255, 255, 255, 0.05);
        animation: slideInUp 0.5s cubic-bezier(0.4, 0, 0.2, 1);
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
        position: relative;
        overflow: hidden;
    }
    
    .chat-message::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        width: 5px;
        height: 100%;
        background: var(--gradient-primary);
        box-shadow: 0 0 20px rgba(99, 102, 241, 0.6);
    }
    
    .chat-message:hover {
        box-shadow: var(--shadow-md), 0 0 30px rgba(99, 102, 241, 0.3);
        transform: translateX(8px) translateY(-2px);
        border-color: rgba(99, 102, 241, 0.4);
    }
    
    @keyframes slideInUp {
        from {
            opacity: 0;
            transform: translateY(30px) scale(0.95);
        }
        to {
            opacity: 1;
            transform: translateY(0) scale(1);
        }
    }
    
    .doctor-message {
        background: linear-gradient(135deg, rgba(99, 102, 241, 0.2) 0%, rgba(139, 92, 246, 0.15) 100%);
        border-left-color: var(--primary);
        border: 1px solid rgba(99, 102, 241, 0.4);
        box-shadow: var(--shadow-sm), 0 0 20px rgba(99, 102, 241, 0.2);
    }
    
    .ai-message {
        background: linear-gradient(135deg, rgba(168, 85, 247, 0.2) 0%, rgba(236, 72, 153, 0.15) 100%);
        border-left-color: var(--accent-purple);
        border: 1px solid rgba(168, 85, 247, 0.4);
        box-shadow: var(--shadow-sm), 0 0 20px rgba(168, 85, 247, 0.2);
    }
    
    .ai-message::before {
        background: linear-gradient(135deg, #a855f7 0%, #ec4899 100%);
        box-shadow: 0 0 20px rgba(168, 85, 247, 0.6);
    }
    
    /* Ultra Modern Button Styles with Glow Effects */
    .stButton > button {
        background: var(--gradient-primary) !important;
        color: white !important;
        border: none !important;
        border-radius: 16px !important;
        padding: 1rem 2rem !important;
        font-weight: 700 !important;
        font-size: 1rem !important;
        box-shadow: var(--shadow-md), var(--shadow-glow) !important;
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1) !important;
        text-transform: none !important;
        letter-spacing: 0.02em !important;
        position: relative !important;
        overflow: hidden !important;
        backdrop-filter: blur(10px) !important;
    }
    
    .stButton > button::before {
        content: '';
        position: absolute;
        top: 50%;
        left: 50%;
        width: 0;
        height: 0;
        border-radius: 50%;
        background: rgba(255, 255, 255, 0.3);
        transform: translate(-50%, -50%);
        transition: width 0.6s, height 0.6s;
    }
    
    .stButton > button:hover::before {
        width: 300px;
        height: 300px;
    }
    
    .stButton > button:hover {
        background: var(--gradient-primary) !important;
        box-shadow: var(--shadow-lg), var(--shadow-glow), 0 0 40px rgba(99, 102, 241, 0.5) !important;
        transform: translateY(-4px) scale(1.05) !important;
        border: 1px solid rgba(255, 255, 255, 0.2) !important;
    }
    
    .stButton > button:active {
        transform: translateY(-2px) scale(1.02) !important;
        box-shadow: var(--shadow-md), var(--shadow-glow) !important;
    }
    
    .stButton > button span {
        position: relative;
        z-index: 1;
    }
    
    /* Ultra Modern Input Styles with Glassmorphism */
    .stTextArea textarea, .stTextInput input {
        border-radius: 16px !important;
        border: 2px solid rgba(255, 255, 255, 0.1) !important;
        padding: 1rem 1.25rem !important;
        background: rgba(15, 23, 42, 0.6) !important;
        backdrop-filter: blur(15px) saturate(180%) !important;
        color: var(--text-primary) !important;
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1) !important;
        font-size: 1rem !important;
        box-shadow: var(--shadow-sm), inset 0 1px 0 rgba(255, 255, 255, 0.05) !important;
    }
    
    .stTextArea textarea:focus, .stTextInput input:focus {
        border-color: var(--primary) !important;
        box-shadow: 0 0 0 4px rgba(99, 102, 241, 0.2), var(--shadow-md), 0 0 30px rgba(99, 102, 241, 0.3) !important;
        outline: none !important;
        background: rgba(15, 23, 42, 0.8) !important;
        transform: scale(1.01) !important;
    }
    
    .stTextArea textarea::placeholder, .stTextInput input::placeholder {
        color: var(--text-muted) !important;
        opacity: 0.7 !important;
    }
    
    /* Empty State */
    .empty-state {
        width: 100%;
        text-align: center;
        border: 2px dashed var(--border-color);
        border-radius: 20px;
        padding: 4rem 2rem;
        background: var(--bg-card);
        color: var(--text-primary);
    }
    
    .empty-state .icon {
        font-size: 4rem;
        margin-bottom: 1.5rem;
        opacity: 0.7;
        filter: drop-shadow(0 4px 8px rgba(99, 102, 241, 0.3));
    }
    
    .empty-state h3 {
        margin: 0 0 0.75rem 0;
        color: var(--text-primary);
        font-size: 1.5rem;
        font-weight: 700;
    }
    
    .empty-state p {
        margin: 0;
        color: var(--text-muted);
        font-size: 1rem;
    }
    
    /* Ultra Modern Chat Container - No scrolling, flows with main page */
    .chat-container {
        background: rgba(20, 27, 45, 0.6);
        backdrop-filter: blur(20px) saturate(180%);
        padding: 2rem;
        border-radius: 24px;
        border: 1px solid rgba(255, 255, 255, 0.1);
        box-shadow: var(--shadow-md), inset 0 1px 0 rgba(255, 255, 255, 0.05);
        /* No max-height or overflow - content flows naturally with main page */
        overflow: visible;
    }
    
    /* Summary Card */
    .summary-card {
        background: var(--gradient-card);
        border: 1px solid var(--border-color);
        margin-top: 1.5rem !important;
        margin-bottom: 2rem !important;
    }
    
    /* Discharge Summary Section - Professional Spacing */
    .summary-card .card-header {
        margin-top: 0 !important;
        margin-bottom: 1.5rem !important;
    }
    
    /* Text Area Spacing */
    .stTextArea {
        margin-top: 1.5rem !important;
        margin-bottom: 2rem !important;
    }
    
    /* Button Groups - Professional Alignment */
    .element-container:has(button) {
        margin-top: 1rem !important;
        margin-bottom: 1rem !important;
    }
    
    /* Summary Action Buttons - Better Spacing */
    [data-testid="column"]:has(button[kind="formSubmit"]) {
        margin-top: 1.5rem !important;
        margin-bottom: 1.5rem !important;
    }
    
    /* Download Buttons Section - Professional Spacing */
    .stDownloadButton {
        margin-top: 1rem !important;
        margin-bottom: 1rem !important;
    }
    
    /* Divider/HR Spacing */
    hr {
        margin-top: 2.5rem !important;
        margin-bottom: 2.5rem !important;
        border-color: rgba(255, 255, 255, 0.1) !important;
        opacity: 0.5 !important;
    }
    
    /* RAG Feedback Section Spacing */
    .stMarkdown h3 {
        margin-top: 2rem !important;
        margin-bottom: 1.5rem !important;
    }
    
    /* Action Buttons Container - Professional Spacing */
    [data-testid="column"] button {
        margin-top: 0.5rem !important;
        margin-bottom: 0.5rem !important;
    }
    
    /* Column spacing for better alignment */
    [data-testid="column"] {
        padding-left: 0.75rem !important;
        padding-right: 0.75rem !important;
    }
    
    /* Text area wrapper spacing */
    .stTextArea > label {
        margin-bottom: 0.75rem !important;
    }
    
    /* Summary section overall spacing */
    [data-testid="column"] .summary-card {
        margin-top: 1rem !important;
        margin-bottom: 2rem !important;
    }
    
    /* Better spacing between sections */
    .main [data-testid="column"] {
        margin-top: 1rem !important;
    }
    
    /* Spacing for button columns */
    [data-testid="column"]:has(button) {
        margin-top: 1rem !important;
        margin-bottom: 1rem !important;
    }
    
    /* Spinner */
    .stSpinner > div {
        border-color: var(--primary) transparent transparent transparent;
    }
    
    /* Loading Animation */
    @keyframes pulse {
        0%, 100% {
            opacity: 1;
        }
        50% {
            opacity: 0.6;
        }
    }
    
    .loading {
        animation: pulse 2s cubic-bezier(0.4, 0, 0.6, 1) infinite;
    }
    
    /* Ultra Modern Card Header with Animated Icon */
    .card-header {
        display: flex;
        align-items: center;
        gap: 1.5rem;
        margin-top: 2.5rem !important;
        margin-bottom: 2rem !important;
        padding-bottom: 1.5rem;
        border-bottom: 2px solid rgba(255, 255, 255, 0.1);
        position: relative;
    }
    
    /* First section header - less top margin */
    .main .element-container:first-of-type .card-header,
    .main [data-testid="column"]:first-of-type .card-header {
        margin-top: 1rem !important;
    }
    
    /* Section headers spacing - add more vertical space */
    .main [data-testid="column"] .card-header {
        margin-top: 2.5rem !important;
        margin-bottom: 2.5rem !important;
    }
    
    .card-header::after {
        content: '';
        position: absolute;
        bottom: -2px;
        left: 0;
        width: 100px;
        height: 2px;
        background: var(--gradient-primary);
        box-shadow: 0 0 20px rgba(99, 102, 241, 0.6);
        animation: slideRight 3s ease-in-out infinite;
    }
    
    @keyframes slideRight {
        0%, 100% { left: 0; }
        50% { left: calc(100% - 100px); }
    }
    
    .card-header-icon {
        width: 56px;
        height: 56px;
        border-radius: 16px;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 1.75rem;
        background: var(--gradient-primary);
        color: white;
        box-shadow: var(--shadow-md), var(--shadow-glow);
        position: relative;
        overflow: hidden;
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
    }
    
    .card-header-icon::before {
        content: '';
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: radial-gradient(circle, rgba(255, 255, 255, 0.3) 0%, transparent 70%);
        animation: iconRotate 3s linear infinite;
    }
    
    @keyframes iconRotate {
        from { transform: rotate(0deg); }
        to { transform: rotate(360deg); }
    }
    
    .card-header-icon:hover {
        transform: scale(1.1) rotate(5deg);
        box-shadow: var(--shadow-lg), var(--shadow-glow), 0 0 40px rgba(99, 102, 241, 0.5);
    }
    
    /* Ultra Modern Sidebar with Glassmorphism */
    [data-testid="stSidebar"] {
        background: rgba(20, 27, 45, 0.8) !important;
        backdrop-filter: blur(20px) saturate(180%) !important;
        border-right: 1px solid rgba(255, 255, 255, 0.1) !important;
        box-shadow: 4px 0 24px rgba(0, 0, 0, 0.3) !important;
    }
    
    [data-testid="stSidebar"] .css-1d391kg {
        background: transparent !important;
    }
    
    [data-testid="stSidebar"] .stMarkdown h1,
    [data-testid="stSidebar"] .stMarkdown h2,
    [data-testid="stSidebar"] .stMarkdown h3 {
        background: var(--gradient-primary) !important;
        -webkit-background-clip: text !important;
        -webkit-text-fill-color: transparent !important;
        background-clip: text !important;
    }
    
    /* Progress Bar */
    .stProgress > div > div > div {
        background: var(--gradient-primary);
    }
    
    /* Success/Error Messages */
    .stSuccess {
        background: rgba(16, 185, 129, 0.1);
        border-left: 4px solid var(--success);
        border-radius: 8px;
        border: 1px solid rgba(16, 185, 129, 0.3);
    }
    
    .stError {
        background: rgba(239, 68, 68, 0.1);
        border-left: 4px solid var(--danger);
        border-radius: 8px;
        border: 1px solid rgba(239, 68, 68, 0.3);
    }
    
    /* Text Area for Summary */
    .stTextArea textarea {
        background: var(--bg-input) !important;
        color: var(--text-primary) !important;
    }
    
    /* Selectbox and other inputs */
    .stSelectbox > div > div {
        background: var(--bg-input);
        border-color: var(--border-color);
    }
    
    /* Markdown text */
    .stMarkdown {
        color: var(--text-primary);
    }
    
    /* Info boxes */
    .stInfo {
        background: rgba(6, 182, 212, 0.1);
        border-left: 4px solid var(--accent-cyan);
        border-radius: 8px;
        border: 1px solid rgba(6, 182, 212, 0.3);
    }
    
    /* Headers */
    h1, h2, h3, h4, h5, h6 {
        color: var(--text-primary) !important;
    }
    
    /* Sidebar text */
    [data-testid="stSidebar"] p, 
    [data-testid="stSidebar"] label,
    [data-testid="stSidebar"] .stMarkdown {
        color: var(--text-primary) !important;
    }
    
    /* Main content text */
    .main .stMarkdown p,
    .main .stMarkdown li {
        color: var(--text-secondary) !important;
    }
    
    /* Expander */
    .streamlit-expanderHeader {
        background: var(--bg-card) !important;
        color: var(--text-primary) !important;
        border: 1px solid var(--border-color) !important;
        border-radius: 8px !important;
    }
    
    .streamlit-expanderContent {
        background: var(--bg-secondary) !important;
        color: var(--text-primary) !important;
    }
    
    /* File uploader */
    .stFileUploader > div > div {
        background: var(--bg-input) !important;
        border-color: var(--border-color) !important;
    }
    
    /* Form */
    .stForm {
        border: 1px solid var(--border-color) !important;
        border-radius: 12px !important;
        padding: 1rem !important;
        background: var(--bg-card) !important;
    }
    
    /* Streamlit Header Bar */
    [data-testid="stHeader"] {
        background: var(--bg-secondary) !important;
        border-bottom: 1px solid var(--border-color) !important;
    }
    
    [data-testid="stHeader"] > div {
        background: var(--bg-secondary) !important;
    }
    
    /* Form Submit Buttons */
    .stForm button[type="submit"],
    .stForm button[kind="formSubmit"],
    form button[type="submit"] {
        background: var(--gradient-primary) !important;
        color: white !important;
        border: none !important;
        border-radius: 12px !important;
        padding: 0.875rem 1.75rem !important;
        font-weight: 600 !important;
        font-size: 0.95rem !important;
        box-shadow: var(--shadow-md) !important;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
        width: 100% !important;
    }
    
    .stForm button[type="submit"]:hover,
    .stForm button[kind="formSubmit"]:hover,
    form button[type="submit"]:hover {
        background: var(--primary-hover) !important;
        box-shadow: var(--shadow-lg), var(--shadow-glow) !important;
        transform: translateY(-2px) !important;
    }
    
    /* All buttons - ensure they have the theme by default */
    button[kind="secondary"],
    button[kind="primary"],
    .stButton > button,
    button {
        background: var(--gradient-primary) !important;
        color: white !important;
        border: none !important;
        border-radius: 12px !important;
        padding: 0.875rem 1.75rem !important;
        font-weight: 600 !important;
        font-size: 0.95rem !important;
        box-shadow: var(--shadow-md) !important;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
    }
    
    button[kind="secondary"]:hover,
    button[kind="primary"]:hover,
    .stButton > button:hover,
    button:hover {
        background: var(--primary-hover) !important;
        box-shadow: var(--shadow-lg), var(--shadow-glow) !important;
        transform: translateY(-2px) !important;
    }
    
    button[kind="secondary"]:active,
    button[kind="primary"]:active,
    .stButton > button:active,
    button:active {
        transform: translateY(0) !important;
    }
    
    /* Modal Styles */
    .modal-overlay {
        position: fixed;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: rgba(0, 0, 0, 0.7);
        display: flex;
        align-items: center;
        justify-content: center;
        z-index: 9999;
        backdrop-filter: blur(5px);
    }
    
    .modal-content {
        background: var(--bg-card);
        border: 1px solid var(--border-color);
        border-radius: 24px;
        padding: 2rem;
        max-width: 800px;
        /* Removed max-height and overflow-y to prevent nested scrolling */
        box-shadow: var(--shadow-lg);
        color: var(--text-primary);
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
    }
    
    .modal-header {
        display: flex;
        justify-content: space-between;
        align-items: center;
        margin-bottom: 1.5rem;
        padding-bottom: 1rem;
        border-bottom: 2px solid var(--border-color);
    }
    
    .modal-header h2 {
        margin: 0;
        color: var(--text-primary);
        font-size: 1.75rem;
        font-weight: 700;
    }
    
    .modal-close {
        background: var(--danger);
        color: white;
        border: none;
        border-radius: 8px;
        padding: 0.5rem 1rem;
        cursor: pointer;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    
    .modal-close:hover {
        background: #dc2626;
        transform: scale(1.05);
    }
    
    .modal-section {
        margin: 1.5rem 0;
        padding: 1rem;
        background: var(--bg-secondary);
        border-radius: 12px;
        border-left: 4px solid var(--primary);
    }
    
    .modal-section h3 {
        margin: 0 0 0.75rem 0;
        color: var(--primary-light);
        font-size: 1.1rem;
        font-weight: 600;
    }
    
    .modal-section p {
        margin: 0.5rem 0;
        color: var(--text-secondary);
        line-height: 1.6;
    }
    
    .modal-section strong {
        color: var(--text-primary);
        font-weight: 600;
    }
    
    /* Fix Streamlit Footer - Completely Remove from Layout */
    footer[data-testid="stFooter"],
    footer,
    .stFooter,
    [data-testid="stFooter"],
    div[data-testid="stFooter"],
    .stApp footer,
    footer.stFooter {
        display: none !important;
        visibility: hidden !important;
        height: 0 !important;
        min-height: 0 !important;
        max-height: 0 !important;
        padding: 0 !important;
        margin: 0 !important;
        border: none !important;
        position: absolute !important;
        top: -9999px !important;
        left: -9999px !important;
        width: 0 !important;
        overflow: hidden !important;
        z-index: -9999 !important;
        opacity: 0 !important;
    }
    
    /* Remove footer from document flow completely */
    .stApp > footer,
    [data-testid="stAppViewContainer"] > footer,
    body > footer {
        display: none !important;
        height: 0 !important;
        margin: 0 !important;
        padding: 0 !important;
    }
    
    /* Ensure content has proper bottom padding so nothing is hidden */
    .main .block-container {
        padding-bottom: 6rem !important;
        margin-bottom: 4rem !important;
    }
    
    /* Fix any fixed position elements that might be blocking content */
    [data-testid="stAppViewContainer"] {
        padding-bottom: 4rem !important;
        margin-bottom: 2rem !important;
    }
    
    /* Ensure main content area has enough space */
    .main {
        padding-bottom: 4rem !important;
        margin-bottom: 2rem !important;
    }
    
    /* Remove any fixed/sticky footer elements */
    *[style*="position: fixed"][style*="bottom"],
    *[style*="position:fixed"][style*="bottom"] {
        bottom: -9999px !important;
    }
    
    /* Ensure body and html don't have footer space */
    body, html {
        padding-bottom: 0 !important;
        margin-bottom: 0 !important;
    }
</style>
<script>
    // Aggressively remove footer on page load and continuously
    function removeFooter() {
        const footers = document.querySelectorAll('footer, [data-testid="stFooter"], .stFooter, footer[data-testid="stFooter"]');
        footers.forEach(footer => {
            footer.style.display = 'none';
            footer.style.visibility = 'hidden';
            footer.style.height = '0';
            footer.style.padding = '0';
            footer.style.margin = '0';
            footer.style.position = 'absolute';
            footer.style.top = '-9999px';
            footer.style.opacity = '0';
            footer.remove(); // Actually remove from DOM
        });
    }
    
    // Run immediately
    removeFooter();
    
    // Run on DOMContentLoaded
    if (document.readyState === 'loading') {
        document.addEventListener('DOMContentLoaded', removeFooter);
    } else {
        removeFooter();
    }
    
    // Run after a short delay to catch dynamically added footers
    setTimeout(removeFooter, 100);
    setTimeout(removeFooter, 500);
    setTimeout(removeFooter, 1000);
    
    // Use MutationObserver to catch any footer that gets added later
    const observer = new MutationObserver(function(mutations) {
        removeFooter();
    });
    
    observer.observe(document.body, {
        childList: true,
        subtree: true
    });
</script>
""", unsafe_allow_html=True)

# Initialize session state
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'current_patient' not in st.session_state:
    st.session_state.current_patient = None
if 'discharge_summary' not in st.session_state:
    st.session_state.discharge_summary = None
if 'autogen_agent' not in st.session_state:
    st.session_state.autogen_agent = None
if 'show_patient_overview' not in st.session_state:
    st.session_state.show_patient_overview = False
if 'selected_similar_patient' not in st.session_state:
    st.session_state.selected_similar_patient = None

class MedicalRAGSystem:
    def __init__(self):
        self.mongo_uri = "mongodb+srv://ishaanroopesh0102:6eShFuC0pNnFFNGm@cluster0.biujjg4.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0"
        self.chroma_path = "vector_db/chroma"
        # Groq API Configuration
        self.groq_api_key = os.getenv("GROQ_API_KEY", "")
        self.groq_base_url = "https://api.groq.com/openai/v1"
        self.groq_model = "meta-llama/llama-4-maverick-17b-128e-instruct"
        self.num_results = 3
        self.http = _http_session()
        self.embedding_cache = _embedding_cache
        
        # Initialize models
        self._load_models()
        self._connect_databases()

    def extract_template_outline(self, template_bytes: bytes) -> list[str]:
        """Extract an ordered list of section headings from a PDF template.

        Heuristic: collect lines that look like headings (short, Title/ALLCAPS, or end with ':').
        Returns a de-duplicated ordered list.
        """
        try:
            reader = PdfReader(BytesIO(template_bytes))
            text = []
            # Inspect first 3 pages for headings
            for i, page in enumerate(reader.pages[:3]):
                try:
                    text.append(page.extract_text() or "")
                except Exception:
                    continue
            joined = "\n".join(text)
            lines = [l.strip() for l in joined.splitlines()]
            candidates: list[str] = []
            for line in lines:
                if not line:
                    continue
                if len(line) < 3 or len(line) > 80:
                    continue
                # Skip page numbers and dates
                if line.lower().startswith("page "):
                    continue
                # Heuristic rules
                looks_like_heading = (
                    line.endswith(":") or
                    (line.isupper() and any(c.isalpha() for c in line)) or
                    (line.istitle() and sum(ch.isalpha() for ch in line) >= 6)
                )
                if looks_like_heading:
                    normalized = line.rstrip(":").strip()
                    if normalized not in candidates:
                        candidates.append(normalized)
            # Keep a reasonable number
            return candidates[:30] if candidates else []
        except Exception:
            return []

    def generate_discharge_summary_with_template(self, patient_data: str, outline_sections: list[str]) -> str:
        """Generate discharge summary following the provided ordered outline sections."""
        outline_bullets = "\n".join([f"- {s}" for s in outline_sections])
        system_prompt = f"""You are an expert medical AI assistant that generates a clinically accurate discharge summary.
Follow the section order EXACTLY as specified by the provided outline. Do not add extra sections; if information is missing, write "[Information not available]".

REQUIRED SECTION ORDER (USE EXACT TITLES):
{outline_bullets}

FORMATTING REQUIREMENTS:
- Use clear section headings on separate lines (e.g., "Name:", "Unit No:", etc.)
- Put each section on its own line with proper spacing
- Use line breaks between major sections
- For lists (medications, diagnoses), put each item on a new line or separate with commas clearly
- Use concise, professional medical language
- Base content solely on the input patient data
- Preserve patient identifiers verbatim if present
- Be brief and factual
- Do NOT use markdown formatting (no asterisks, no bold, no headers). Use plain text with line breaks for structure."""

        user_prompt = f"""Generate a discharge summary STRICTLY following the section list above, based only on this data:\n\n{patient_data}\n\nReturn plain text with the exact section headings in order."""

        if not self.groq_api_key:
            return "‚ùå GROQ_API_KEY environment variable is not set. Please set it in your .env file."

        try:
            response = self.http.post(
                f"{self.groq_base_url}/chat/completions",
                json={
                    "model": self.groq_model,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    "temperature": 0.3,
                        "top_p": 0.85,
                    "max_tokens": 4000
                },
                headers={
                    "Authorization": f"Bearer {self.groq_api_key}",
                    "Content-Type": "application/json"
                },
                timeout=120.0
            )
            if response.ok:
                data = response.json()
                if "choices" in data and len(data["choices"]) > 0:
                    content = data["choices"][0].get("message", {}).get("content", "")
                    return content.strip() if content.strip() else "Summary generated successfully."
                else:
                    return "‚ùå Error: Invalid response format from Groq API"
            else:
                error_msg = response.text
                if response.status_code == 401:
                    return "‚ùå Error: Invalid Groq API key. Please check your GROQ_API_KEY."
                elif response.status_code == 429:
                    return "‚ùå Error: Groq API rate limit exceeded. Please wait 30-60 seconds before trying again. Consider upgrading your Groq API plan for higher rate limits."
                else:
                    return f"‚ùå Error generating summary (HTTP {response.status_code}): {error_msg}"
        except requests.exceptions.Timeout:
            return "‚ùå Error: Request to Groq API timed out. Please try again."
        except Exception as e:
            return f"‚ùå Error connecting to Groq API: {str(e)}"

    def generate_pdf_from_text(self, text: str, template_bytes: bytes | None = None) -> bytes:
        """Generate a PDF from plain text. If a PDF template is provided, overlay text pages on template pages.

        Args:
            text: The discharge summary text
            template_bytes: Raw bytes of a PDF template (optional)
        Returns:
            PDF file bytes
        """
        # Determine page size: use template first page if provided, else A4
        page_size = A4
        template_reader = None
        if template_bytes:
            try:
                template_reader = PdfReader(BytesIO(template_bytes))
                first_page = template_reader.pages[0]
                width = float(first_page.mediabox.width)
                height = float(first_page.mediabox.height)
                page_size = (width, height)
            except Exception:
                template_reader = None

        # Build a PDF with flowing text
        buf = BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=page_size, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()
        body_style = ParagraphStyle(
            name="Body",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            alignment=TA_LEFT,
        )
        story = []
        for para in text.split("\n\n"):
            story.append(Paragraph(para.replace("\n", "<br/>"), body_style))
            story.append(Spacer(1, 0.18 * inch))
        doc.build(story)
        generated_pdf_bytes = buf.getvalue()

        # If no template, return generated bytes directly
        if not template_reader:
            return generated_pdf_bytes

        # Merge generated pages onto template pages
        gen_reader = PdfReader(BytesIO(generated_pdf_bytes))
        writer = PdfWriter()
        num_pages = max(len(template_reader.pages), len(gen_reader.pages))

        for i in range(num_pages):
            template_page = None
            if i < len(template_reader.pages):
                template_page = template_reader.pages[i]

            if i < len(gen_reader.pages):
                gen_page = gen_reader.pages[i]
                if template_page is not None:
                    try:
                        # Overlay generated content on top of template background
                        template_page.merge_page(gen_page)
                        writer.add_page(template_page)
                    except Exception:
                        # Fallback to adding generated page if merge fails
                        writer.add_page(gen_page)
                else:
                    writer.add_page(gen_page)
            else:
                # No generated content for this page, keep template page
                if template_page is not None:
                    writer.add_page(template_page)

        out_buf = BytesIO()
        writer.write(out_buf)
        return out_buf.getvalue()

    def generate_docx_from_text(self, text: str) -> bytes:
        """Generate a DOCX file from plain text, preserving paragraphs and line breaks."""
        doc = Document()
        # Add title if the first line looks like a heading
        lines = text.split("\n")
        if lines and len(lines[0]) <= 120 and any(ch.isalpha() for ch in lines[0]):
            doc.add_heading(lines[0].strip(), level=1)
            text = "\n".join(lines[1:])
        for block in text.split("\n\n"):
            for ln in block.split("\n"):
                doc.add_paragraph(ln)
            doc.add_paragraph("")
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
    
    def _load_models(self):
        """Load Bio ClinicalBERT model for embeddings"""
        with st.spinner("Loading Bio ClinicalBERT model..."):
            self.tokenizer, self.model = _load_tokenizer_model()
    
    def _connect_databases(self):
        """Connect to MongoDB and ChromaDB"""
        try:
            # MongoDB connection
            self.mongo_client = _connect_mongo(self.mongo_uri)
            self.db = self.mongo_client["hospital_db"]
            self.patients_collection = self.db["test_patients"]
            
            # ChromaDB connection
            self.chroma_client, self.chroma_collection = _connect_chroma(self.chroma_path)
            
            st.success("‚úÖ Connected to databases successfully")
        except Exception as e:
            st.error(f"‚ùå Database connection failed: {str(e)}")
    
    def embed_text(self, text: str) -> List[float]:
        """Generate embedding for text using Bio ClinicalBERT with caching"""
        # Check if tokenizer/model is available
        if self.tokenizer is None or self.model is None:
            raise RuntimeError("Transformers model not available. Cannot generate embeddings. Please ensure transformers library is installed and models are loaded.")
        
        # Check cache first
        text_hash = _get_text_hash(text)
        if text_hash in self.embedding_cache:
            return self.embedding_cache[text_hash]
        
        # Generate embedding
        with torch.no_grad():
            inputs = self.tokenizer(text, return_tensors="pt", truncation=True, max_length=256)
            if torch.cuda.is_available():
                inputs = {k: v.to("cuda") for k, v in inputs.items()}
            outputs = self.model(**inputs)
            cls_embedding = outputs.last_hidden_state[:, 0, :]
            emb = cls_embedding.squeeze(0)
            if emb.is_cuda:
                emb = emb.to("cpu")
            embedding = emb.tolist()
        
        # Cache the embedding
        self.embedding_cache[text_hash] = embedding
        return embedding
    
    def format_patient_fields(self, record: Dict) -> str:
        """Format patient record fields for embedding and AI context"""
        fields = [
            "name", "unit no", "admission date", "date of birth", "sex", "service",
            "allergies", "attending", "chief complaint", "major surgical or invasive procedure",
            "history of present illness", "past medical history", "social history",
            "family history", "physical exam", "pertinent results", "medications on admission",
            "brief hospital course", "discharge medications", "discharge diagnosis",
            "discharge condition", "discharge instructions", "follow-up", "discharge disposition"
        ]
        parts = [f"{field.title()}: {record.get(field, '')}" for field in fields if record.get(field)]
        
        # If this is a preprocessed patient with full document, include it for complete context
        if record.get('_full_document'):
            parts.append(f"\n\nFULL PATIENT DOCUMENT:\n{record.get('_full_document')}")
        
        # If discharge summary is available, include it
        if record.get('_discharge_summary'):
            parts.append(f"\n\nDISCHARGE SUMMARY:\n{record.get('_discharge_summary')}")
        
        return " ".join(parts) if parts else "No patient information available"
    
    def get_patient_by_unit_no(self, unit_no: str) -> Optional[Dict]:
        """Retrieve patient record from MongoDB"""
        try:
            record = self.patients_collection.find_one({"unit no": int(unit_no)})
            return record
        except Exception as e:
            st.error(f"Error retrieving patient: {str(e)}")
            return None
    
    def get_all_patients_list(self) -> List[Dict]:
        """Get list of all patients with name and unit number"""
        try:
            # Query MongoDB for all patients, only get name and unit no
            patients = list(self.patients_collection.find(
                {},
                {"name": 1, "unit no": 1, "_id": 0}  # Only return name and unit no
            ))
            
            # Format patient data
            patient_list = []
            for patient in patients:
                name = patient.get("name", "Unknown")
                unit_no = patient.get("unit no", "N/A")
                if name and name != "Unknown" and unit_no and unit_no != "N/A":
                    patient_list.append({
                        "name": str(name),
                        "unit_no": str(unit_no),
                        "display": f"{name} (Unit {unit_no})"
                    })
            
            # Sort by name
            patient_list.sort(key=lambda x: x["name"])
            return patient_list
        except Exception as e:
            st.error(f"Error retrieving patients list: {str(e)}")
            return []
    
    def search_similar_cases(self, query_text: str, n_results: int = 3) -> List[Dict]:
        """Search for similar cases using RAG - excludes MongoDB patients"""
        try:
            # Get list of MongoDB patient unit numbers to exclude
            mongo_unit_nos = set()
            try:
                mongo_patients = list(self.patients_collection.find({}, {"unit no": 1, "_id": 0}))
                mongo_unit_nos = {str(patient.get("unit no", "")) for patient in mongo_patients if patient.get("unit no")}
            except Exception as e:
                st.warning(f"Could not fetch MongoDB patients for filtering: {str(e)}")
            
            # Search with more results to account for filtering
            query_embedding = self.embed_text(query_text)
            results = self.chroma_collection.query(
                query_embeddings=[query_embedding],
                n_results=n_results * 3,  # Get more results to filter
                include=["documents", "metadatas", "distances"]
            )
            
            similar_cases = []
            for i in range(len(results["documents"][0])):
                metadata = results["metadatas"][0][i] if i < len(results["metadatas"][0]) else {}
                unit_no = str(metadata.get('unit_no') or metadata.get('unit no', ''))
                
                # Skip if this patient exists in MongoDB (only show ChromaDB/preprocessed patients)
                if unit_no and unit_no in mongo_unit_nos:
                    continue
                
                similar_cases.append({
                    "document": results["documents"][0][i],
                    "metadata": metadata,
                    "similarity": 1 - results["distances"][0][i] if i < len(results["distances"][0]) else 0.0
                })
                
                # Stop when we have enough results
                if len(similar_cases) >= n_results:
                    break
            
            return similar_cases
        except Exception as e:
            st.error(f"Error searching similar cases: {str(e)}")
            return []
    
    def _clean_markdown_formatting(self, text: str) -> str:
        """Remove markdown formatting like asterisks, bold markers, etc."""
        import re
        # Remove bold/italic markers (**text**, *text*)
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Remove **bold**
        text = re.sub(r'\*([^*]+)\*', r'\1', text)  # Remove *italic*
        # Remove standalone asterisks used as bullets
        text = re.sub(r'^\s*\*\s+', '', text, flags=re.MULTILINE)
        # Remove markdown headers (# Header)
        text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)
        return text
    
    def generate_discharge_summary(self, patient_data: str, similar_cases: List[Dict] = None) -> str:
        """Generate discharge summary using Groq API"""
        if not self.groq_api_key:
            return "‚ùå GROQ_API_KEY environment variable is not set. Please set it in your .env file."
        
        system_prompt = """You are an expert medical AI assistant that generates structured, clinically accurate discharge summaries.
Base your summary entirely on the INPUT PATIENT DATA provided.
The discharge summary MUST include: Name, Unit No, Date Of Birth, Sex, Admission/Discharge Dates, Attending, Chief Complaint, Procedure, History, Physical Exam (on Admission), Pertinent Results, Brief Hospital Course, Medications on Admission, Discharge Medications, Discharge Instructions, Discharge Disposition, Discharge Diagnosis, Discharge Condition, Follow-up.

FORMATTING REQUIREMENTS:
- Use clear section headings on separate lines (e.g., "Name:", "Unit No:", "Admission Date:", etc.)
- Put each section on its own line with proper spacing
- Use line breaks between major sections
- For Name, Unit No, Date of Birth, and Sex, put them on separate lines at the top
- For lists (medications, diagnoses), put each item on a new line or separate with commas clearly
- If information is missing, state "[Information not available]"
- Use concise, professional medical language. Be brief and factual.
- Do NOT use markdown formatting (no asterisks, no bold, no headers). Use plain text with line breaks for structure.

EXAMPLE FORMAT:
Name: [Name]
Unit No: [Unit No]
Date of Birth: [DOB]
Sex: [Sex]

Admission Date: [Date]
Discharge Date: [Date]
Attending: [Name]

Chief Complaint: [Complaint]

Procedure: [Procedure details]

History Of Present Illness: [History]

Past Medical History: [History]

Physical Exam: [Exam details]

[Continue with other sections, each on separate lines]"""

        user_prompt = f"""Generate a discharge summary for this patient:
{patient_data}

Extract Name, Unit No, Date of Birth, and Sex exactly as provided."""

        try:
            response = self.http.post(
                f"{self.groq_base_url}/chat/completions",
                json={
                    "model": self.groq_model,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    "temperature": 0.3,
                        "top_p": 0.85,
                    "max_tokens": 4000
                },
                headers={
                    "Authorization": f"Bearer {self.groq_api_key}",
                    "Content-Type": "application/json"
                },
                timeout=120.0  # 2 minute timeout
            )

            if response.ok:
                data = response.json()
                if "choices" in data and len(data["choices"]) > 0:
                    content = data["choices"][0].get("message", {}).get("content", "")
                    return content.strip() if content.strip() else "Summary generated successfully."
                else:
                    return "‚ùå Error: Invalid response format from Groq API"
            else:
                error_msg = response.text
                if response.status_code == 401:
                    return "‚ùå Error: Invalid Groq API key. Please check your GROQ_API_KEY."
                elif response.status_code == 429:
                    return "‚ùå Error: Groq API rate limit exceeded. Please wait 30-60 seconds before trying again. Consider upgrading your Groq API plan for higher rate limits."
                else:
                    return f"‚ùå Error generating summary (HTTP {response.status_code}): {error_msg}"
        except requests.exceptions.Timeout:
            return "‚ùå Error: Request to Groq API timed out. Please try again."
        except Exception as e:
            return f"‚ùå Error connecting to Groq API: {str(e)}"

    # --- START: NEW FEEDBACK LOOP METHOD ---
    def add_summary_to_vector_db(self, patient_info: Dict, summary_text: str):
        """
        Embeds the finalized discharge summary and adds it to the ChromaDB collection.
        This serves as the feedback loop, adding a high-quality, human-reviewed
        document back into the RAG system.
        """
        if not summary_text or not patient_info:
            st.warning("No summary text or patient info available to add.")
            return False

        unit_no = patient_info.get('unit no', 'unknown')
        patient_name = patient_info.get('name', 'Unknown')
        
        try:
            # 1. Generate embedding for the new summary
            # Check if transformers model is available
            if self.tokenizer is None or self.model is None:
                st.error("‚ùå Cannot add summary to knowledge base: Embedding model not available. The transformers model needs to be loaded. This feature requires the Bio ClinicalBERT model to be initialized.")
                st.info("üí° Tip: The embedding model loads automatically when FastAPI backend is available. If you're in fallback mode, wait for FastAPI to finish loading (2-3 minutes), then try again.")
                return False
            
            summary_embedding = self.embed_text(summary_text)
            
            # 2. Prepare a unique ID
            # Using unit_no and timestamp allows for multiple summary versions
            doc_id = f"summary_{unit_no}_{int(time.time())}"
            
            # 3. Prepare metadata
            metadata = {
                "unit_no": str(unit_no),
                "name": patient_name,
                "summary": summary_text[:500],  # Store a preview in metadata
                "source_type": "feedback_summary" # Tag this as a human-reviewed entry
            }
            
            # 4. Add to ChromaDB
            self.chroma_collection.add(
                embeddings=[summary_embedding],
                documents=[summary_text],  # Store the full summary as the document
                metadatas=[metadata],
                ids=[doc_id]
            )
            
            # 5. Show notification (as requested)
            # st.toast is available in newer Streamlit; fall back to success if missing
            try:
                st.toast(f"Database updated: Summary for {unit_no} added.", icon="‚úÖ")
            except Exception:
                st.success(f"Database updated: Summary for {unit_no} added.")
            return True
        
        except Exception as e:
            st.error(f"‚ùå Error adding feedback summary to vector DB: {str(e)}")
            st.exception(e) # Print full error
            return False
    # --- END: NEW FEEDBACK LOOP METHOD ---
    
class AutoGenMedicalAgent:
    def __init__(self, rag_system: MedicalRAGSystem, api_client: FastAPIClient = None):
        self.rag_system = rag_system
        self.api_client = api_client
        self.agent = None
        self.user_proxy = None
        self._initialize_agent()
    
    def _initialize_agent(self):
        """Initialize AutoGen medical assistant agent"""
        # Skip AutoGen initialization to avoid API errors
        # Use FastAPI backend instead
        pass
    
    def chat_with_doctor(self, message: str, patient_data: Dict = None) -> str:
        """Handle conversation with doctor - uses FastAPI if available"""
        try:
            # Use FastAPI if available, otherwise fallback
            if self.api_client and self.api_client.health_check():
                return self.api_client.chat(message, patient_data)
            else:
                return self._fallback_chat(message, patient_data)
        except Exception as e:
            return f"‚ùå Error in conversation: {str(e)}"
    
    def _fallback_chat(self, message: str, patient_data: Dict = None) -> str:
        """Fallback chat using direct Groq API interaction - optimized for speed"""
        if not self.rag_system.groq_api_key:
            return "‚ùå GROQ_API_KEY environment variable is not set. Please set it in your .env file."
        
        try:
            # Check if user is asking for discharge summary generation
            if "discharge summary" in message.lower() or "generate summary" in message.lower():
                if patient_data:
                    # Use the existing discharge summary generation method
                    patient_text = self.rag_system.format_patient_fields(patient_data)
                    return self.rag_system.generate_discharge_summary(patient_text)
                else:
                    return "‚ùå Please select a patient first to generate a discharge summary."
            
            # Add patient context with actual patient data if available
            patient_context = ""
            if patient_data:
                # Format patient data for context
                patient_text = self.rag_system.format_patient_fields(patient_data)
                patient_name = patient_data.get('name', 'Unknown')
                unit_no = patient_data.get('unit no', 'N/A')
                patient_context = f"""

CURRENT PATIENT INFORMATION:
{patient_text}

Note: The person you are talking to is a medical professional (doctor or nurse) asking questions about this patient. Answer their questions based on the patient information provided above."""
            
            system_prompt = """You are a medical AI assistant designed to help doctors and nurses with clinical documentation and medical questions. 
The person you are talking to is a medical professional (doctor or nurse), NOT a patient. 
When patient information is provided, use it to answer questions about that specific patient.
IMPORTANT: Only use the patient information provided in the CURRENT PATIENT INFORMATION section below. Do not reference any previous patients or conversations.
Provide accurate, helpful responses based on the patient data. Be concise but thorough."""
            
            user_message = message
            if patient_context:
                # Add explicit instruction to use only the current patient
                user_message = f"{message}\n\n{patient_context}\n\nRemember: Only answer questions about the patient information provided above. Do not reference any other patients."
            else:
                # If no patient data, remind that patient needs to be selected
                user_message = f"{message}\n\nNote: No patient is currently selected. Please select a patient first to get patient-specific information."
            
            # Optimize request for faster response using Groq API with retry logic
            import time
            max_retries = 3
            last_error = None
            
            for attempt in range(max_retries):
                try:
                    response = self.rag_system.http.post(
                        f"{self.rag_system.groq_base_url}/chat/completions",
                        json={
                            "model": self.rag_system.groq_model,
                            "messages": [
                                {"role": "system", "content": system_prompt},
                                {"role": "user", "content": user_message}
                            ],
                            "temperature": 0.4,
                            "top_p": 0.85,
                            "max_tokens": 500  # Increased to allow proper responses
                        },
                        headers={
                            "Authorization": f"Bearer {self.rag_system.groq_api_key}",
                            "Content-Type": "application/json"
                        },
                        timeout=30.0
                    )
                    
                    if response.ok:
                        data = response.json()
                        if "choices" in data and len(data["choices"]) > 0:
                            content = data["choices"][0].get("message", {}).get("content", "")
                            return content.strip() if content.strip() else "I'm here to help with medical questions. How can I assist you?"
                        else:
                            return "‚ùå Error: Invalid response format from Groq API"
                    else:
                        error_msg = response.text
                        if response.status_code == 401:
                            return "‚ùå Error: Invalid Groq API key. Please check your GROQ_API_KEY."
                        elif response.status_code == 429:
                            # Rate limit - wait and retry with exponential backoff
                            if attempt < max_retries - 1:
                                wait_time = min(2 ** (attempt + 2), 30)  # 4s, 8s, 16s (max 30s)
                                time.sleep(wait_time)
                                continue
                            return "‚ùå Error: Groq API rate limit exceeded. Please wait 30-60 seconds before trying again. Consider upgrading your Groq API plan for higher rate limits."
                        else:
                            # For other errors, retry once
                            if attempt < max_retries - 1 and response.status_code >= 500:
                                time.sleep(2 ** attempt)
                                continue
                            return f"‚ùå Error connecting to Groq API (HTTP {response.status_code}): {error_msg}"
                except requests.exceptions.Timeout:
                    if attempt < max_retries - 1:
                        time.sleep(2 ** attempt)
                        continue
                    return "‚è±Ô∏è Request timed out. Please try again."
                except Exception as e:
                    last_error = e
                    if attempt < max_retries - 1:
                        time.sleep(1)
                        continue
            
            # If all retries failed
            if last_error:
                error_str = str(last_error)
                if "429" in error_str or "rate limit" in error_str.lower():
                    return "‚ùå Error: Groq API rate limit exceeded. Please wait 30-60 seconds before trying again."
                return f"‚ùå Error: {error_str}"
            return "‚ùå Error: Failed to connect to Groq API after multiple attempts."
                
        except requests.exceptions.Timeout:
            return "‚è±Ô∏è Request timed out. Please try again with a shorter message."
        except Exception as e:
            return f"‚ùå Error in fallback chat: {str(e)}"

def main():
    # Initialize FastAPI client
    if 'api_client' not in st.session_state:
        st.session_state.api_client = FastAPIClient()
    
    # Check FastAPI availability
    use_fastapi = st.session_state.api_client.health_check()
    if use_fastapi:
        st.session_state.use_fastapi = True
        # Initialize with FastAPI client
        if 'rag_system' not in st.session_state:
            # Still need RAG system for some operations (formatting, PDF generation)
            try:
                st.session_state.rag_system = MedicalRAGSystem()
            except Exception as e:
                st.warning(f"‚ö†Ô∏è Could not initialize full RAG system: {str(e)}. Some features may be limited.")
        # Always initialize autogen_agent if not present
        if 'autogen_agent' not in st.session_state:
            try:
                st.session_state.autogen_agent = AutoGenMedicalAgent(
                    st.session_state.rag_system if 'rag_system' in st.session_state else None,
                    st.session_state.api_client
                )
            except Exception as e:
                st.error(f"‚ùå Failed to initialize AI agent: {str(e)}")
    else:
        st.session_state.use_fastapi = False
        # Initialize RAG system as fallback
        if 'rag_system' not in st.session_state:
            with st.spinner("Initializing Medical RAG System (FastAPI unavailable, using fallback)..."):
                try:
                    st.session_state.rag_system = MedicalRAGSystem()
                    st.session_state.autogen_agent = AutoGenMedicalAgent(st.session_state.rag_system, None)
                    st.warning("""
                    ‚ö†Ô∏è **FastAPI backend not available. Using fallback mode.**
                    
                    **To enable FastAPI for better performance:**
                    1. Open a new terminal/command prompt
                    2. Navigate to this project directory
                    3. Run: `python start_api.py` (or `start_api.bat` on Windows)
                    4. Wait for "Application startup complete" message
                    5. Refresh this page
                    
                    The app will work in fallback mode, but responses will be slower.
                    """)
                except Exception as e:
                    st.error(f"‚ùå Failed to initialize system: {str(e)}")
                    st.stop()
        # Ensure autogen_agent is initialized even in fallback mode
        if 'autogen_agent' not in st.session_state:
            try:
                st.session_state.autogen_agent = AutoGenMedicalAgent(
                    st.session_state.rag_system if 'rag_system' in st.session_state else None,
                    None
                )
            except Exception as e:
                st.error(f"‚ùå Failed to initialize AI agent: {str(e)}")

    # Sidebar preferences and CSS
    with st.sidebar:
        st.header("‚öôÔ∏è Preferences")
        minimal_ui = st.checkbox("Minimal UI", value=st.session_state.get('minimal_ui', False))
        st.session_state.minimal_ui = minimal_ui

        st.markdown("---")
        st.header("üìé Insurance Template")
        template_file = st.file_uploader("Upload PDF template (optional)", type=["pdf"], accept_multiple_files=False)
        if template_file is not None:
            st.session_state["template_pdf_bytes"] = template_file.read()
            # Extract outline from template (module-level helper to avoid class ordering issues)
            outline = extract_template_outline(st.session_state["template_pdf_bytes"])
            if outline:
                st.session_state["template_outline"] = outline
                st.success("Template loaded. Outline detected and will be used for generation.")
                with st.expander("Detected Section Order"):
                    for s in outline:
                        st.write(f"‚Ä¢ {s}")
            else:
                st.session_state.pop("template_outline", None)
                st.warning("Template loaded but no clear section outline was detected. Will generate standard summary.")
        elif "template_pdf_bytes" not in st.session_state:
            st.info("No template uploaded. Summaries will be generated as plain text or basic PDF.")

    # Only load minimal CSS if minimal_ui is enabled, otherwise the modern CSS is already loaded at the top
    if st.session_state.get('minimal_ui', False):
        st.markdown(_get_css(True), unsafe_allow_html=True)

    # Header with modern dark design
    st.markdown("""
    <div class="main-header">
        <h1>üè• Medical Discharge Summary Assistant</h1>
        <p>AI-Powered Clinical Documentation with RAG and AutoGen Integration</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Status indicators with modern cards
    col_status1, col_status2, col_status3, col_status4 = st.columns(4)
    
    with col_status1:
        st.markdown("""
        <div class="metric-card">
            <h4>üü¢ System Ready</h4>
            <p>All systems operational</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col_status2:
        st.markdown("""
        <div class="metric-card">
            <h4>‚ö° Fast Mode</h4>
            <p>Optimized responses</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col_status3:
        st.markdown("""
        <div class="metric-card">
            <h4>üß† AI Active</h4>
            <p>LLaMA 3 + RAG</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col_status4:
        patient_count = "1" if st.session_state.current_patient else "0"
        st.markdown(f"""
        <div class="metric-card">
            <h4>üë§ Patient</h4>
            <p>{patient_count} selected</p>
    </div>
    """, unsafe_allow_html=True)
    
    # RAG system is already initialized above
    
    # Sidebar for patient search
    with st.sidebar:
        st.header("üîç Patient Search")
        
        # Get all patients for dropdown
        @st.cache_data(ttl=300)  # Cache for 5 minutes
        def load_patients_list():
            """Load list of all patients"""
            try:
                if st.session_state.get('use_fastapi', False):
                    return st.session_state.api_client.get_all_patients()
                else:
                    return st.session_state.rag_system.get_all_patients_list()
            except Exception as e:
                st.error(f"Error loading patients: {str(e)}")
                return []
        
        patients_list = load_patients_list()
        
        # Create dropdown options
        if patients_list:
            patient_options = ["-- Select a patient --"] + [p["display"] for p in patients_list]
            patient_dict = {p["display"]: p for p in patients_list}
        else:
            patient_options = ["-- No patients available --"]
            patient_dict = {}
        
        # Patient dropdown selection
        selected_patient_display = st.selectbox(
            "Select Patient",
            options=patient_options,
            index=0,
            key="patient_dropdown"
        )
        
        # Track previous patient to detect changes
        if "previous_patient_unit_no" not in st.session_state:
            st.session_state.previous_patient_unit_no = None
        
        # Handle patient selection from dropdown
        if selected_patient_display and selected_patient_display != "-- Select a patient --" and selected_patient_display != "-- No patients available --":
            selected_patient_info = patient_dict.get(selected_patient_display)
            if selected_patient_info:
                selected_unit_no = selected_patient_info["unit_no"]
                
                # Check if patient changed
                if st.session_state.previous_patient_unit_no != selected_unit_no:
                    with st.spinner("Loading patient..."):
                        try:
                            # Use FastAPI if available
                            if st.session_state.get('use_fastapi', False):
                                patient = st.session_state.api_client.get_patient(selected_unit_no)
                            else:
                                patient = st.session_state.rag_system.get_patient_by_unit_no(selected_unit_no)
                            
                            if patient:
                                # Clear chat history and discharge summary when switching patients
                                st.session_state.chat_history = []
                                st.session_state.discharge_summary = None
                                if "editable_summary" in st.session_state:
                                    st.session_state.editable_summary = None
                                
                                # Set new patient
                                st.session_state.current_patient = patient
                                st.session_state.previous_patient_unit_no = selected_unit_no
                                
                                st.success(f"‚úÖ Loaded: {patient.get('name', 'Unknown')}")
                                st.rerun()  # Rerun to refresh UI
                        except Exception as e:
                            st.error(f"‚ùå Error loading patient: {str(e)}")
        
        st.markdown("---")
        st.markdown("**Or search by Unit Number:**")
        
        # Patient search form (keep existing functionality)
        with st.form("patient_search"):
            unit_no = st.text_input("Unit Number", placeholder="Enter patient unit number")
            search_button = st.form_submit_button("üîç Search Patient")
            
            if search_button and unit_no:
                with st.spinner("Searching for patient..."):
                    try:
                        # Use FastAPI if available
                        if st.session_state.get('use_fastapi', False):
                            patient = st.session_state.api_client.get_patient(unit_no)
                        else:
                            patient = st.session_state.rag_system.get_patient_by_unit_no(unit_no)
                        if patient:
                            # Clear chat history, discharge summary, and similar cases when switching patients
                            if st.session_state.previous_patient_unit_no != str(patient.get('unit no', '')):
                                st.session_state.chat_history = []
                                st.session_state.discharge_summary = None
                                if "editable_summary" in st.session_state:
                                    st.session_state.editable_summary = None
                            
                            st.session_state.current_patient = patient
                            st.session_state.previous_patient_unit_no = str(patient.get('unit no', ''))
                            
                            st.markdown(f"""
                            <div style="background: rgba(16, 185, 129, 0.15); border: 1px solid rgba(16, 185, 129, 0.4); border-radius: 12px; padding: 1rem; margin: 0.5rem 0;">
                                <p style="color: var(--success); margin: 0; font-weight: 600;">‚úÖ Found patient: {patient.get('name', 'Unknown')}</p>
                            </div>
                            """, unsafe_allow_html=True)
                            st.rerun()  # Rerun to refresh UI
                        else:
                            st.markdown("""
                            <div style="background: rgba(239, 68, 68, 0.15); border: 1px solid rgba(239, 68, 68, 0.4); border-radius: 12px; padding: 1rem; margin: 0.5rem 0;">
                                <p style="color: var(--danger); margin: 0; font-weight: 600;">‚ùå Patient not found</p>
                            </div>
                            """, unsafe_allow_html=True)
                    except Exception as e:
                        st.markdown(f"""
                        <div style="background: rgba(239, 68, 68, 0.15); border: 1px solid rgba(239, 68, 68, 0.4); border-radius: 12px; padding: 1rem; margin: 0.5rem 0;">
                            <p style="color: var(--danger); margin: 0; font-weight: 600;">‚ùå Error: {str(e)}</p>
                        </div>
                        """, unsafe_allow_html=True)
        
        # Display current patient info - ensure it updates when patient changes
        if st.session_state.current_patient:
            st.markdown("### üë§ Current Patient")
            patient = st.session_state.current_patient
            
            st.markdown(f"""
            <div class="patient-card">
                <h4 style="color: var(--text-primary); margin-bottom: 1.5rem; font-size: 1.3rem; font-weight: 700;">üìã {patient.get('name', 'Unknown')}</h4>
                <p style="color: var(--text-secondary); margin: 0.75rem 0; font-size: 0.95rem;"><strong style="color: var(--primary-light);">Unit No:</strong> {patient.get('unit no', 'N/A')}</p>
                <p style="color: var(--text-secondary); margin: 0.75rem 0; font-size: 0.95rem;"><strong style="color: var(--primary-light);">DOB:</strong> {patient.get('date of birth', 'N/A')}</p>
                <p style="color: var(--text-secondary); margin: 0.75rem 0; font-size: 0.95rem;"><strong style="color: var(--primary-light);">Sex:</strong> {patient.get('sex', 'N/A')}</p>
                <p style="color: var(--text-secondary); margin: 0.75rem 0; font-size: 0.95rem;"><strong style="color: var(--primary-light);">Service:</strong> {patient.get('service', 'N/A')}</p>
                <p style="color: var(--text-secondary); margin: 0.75rem 0; font-size: 0.95rem;"><strong style="color: var(--primary-light);">Attending:</strong> {patient.get('attending', 'N/A')}</p>
            </div>
            """, unsafe_allow_html=True)
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown("""
        <div class="card-header" style="margin-top: 1rem; margin-bottom: 2.5rem;">
            <div class="card-header-icon">üí¨</div>
            <h3 style="margin: 0; color: var(--text-dark);">AI Medical Assistant</h3>
        </div>
        """, unsafe_allow_html=True)
        
        # Chat interface
        if st.session_state.current_patient:
            # Chat container with better styling
            st.markdown("""
            <div class="chat-container">
            """, unsafe_allow_html=True)
            
            # Display chat history
            if st.session_state.chat_history:
                for message in st.session_state.chat_history:
                    if message["role"] == "doctor":
                        st.markdown(f"""
                        <div class="chat-message doctor-message">
                            <div style="display: flex; align-items: flex-start; gap: 1rem;">
                                <div style="background: var(--gradient-primary); color: white; border-radius: 50%; width: 44px; height: 44px; display: flex; align-items: center; justify-content: center; font-size: 1.3rem; flex-shrink: 0; box-shadow: var(--shadow-md);">üë®‚Äç‚öïÔ∏è</div>
                                <div style="flex: 1;">
                                    <div style="font-weight: 700; color: var(--primary-light); margin-bottom: 0.5rem; font-size: 0.9rem; letter-spacing: 0.02em;">Doctor</div>
                                    <div style="color: var(--text-primary); line-height: 1.7; font-size: 0.95rem;">{message["content"]}</div>
                                </div>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    else:
                        st.markdown(f"""
                        <div class="chat-message ai-message">
                            <div style="display: flex; align-items: flex-start; gap: 1rem;">
                                <div style="background: var(--gradient-primary); color: white; border-radius: 50%; width: 44px; height: 44px; display: flex; align-items: center; justify-content: center; font-size: 1.3rem; flex-shrink: 0; box-shadow: var(--shadow-md);">ü§ñ</div>
                                <div style="flex: 1;">
                                    <div style="font-weight: 700; color: var(--accent-purple); margin-bottom: 0.5rem; font-size: 0.9rem; letter-spacing: 0.02em;">AI Assistant</div>
                                    <div style="color: var(--text-primary); line-height: 1.7; font-size: 0.95rem;">{message["content"]}</div>
                                </div>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
            else:
                st.markdown("""
                <div style="text-align: center; color: var(--text-muted); padding: 3rem 2rem; font-style: italic; font-size: 1.1rem;">
                    üëã Start a conversation with the AI assistant...
                </div>
                """, unsafe_allow_html=True)
            
            st.markdown("</div>", unsafe_allow_html=True)
            
            # Chat input form
            with st.form("chat_form", clear_on_submit=False):
                st.markdown("**<span style='color: var(--text-dark); font-weight: 600;'>Ask the AI assistant:</span>**", unsafe_allow_html=True)
                user_message = st.text_area(
                    "message_input", 
                    placeholder="e.g., Generate a discharge summary for this patient",
                    height=100,
                    label_visibility="collapsed"
                )
                
                col_send, col_clear = st.columns([1, 1])
                with col_send:
                    send_button = st.form_submit_button("üí¨ Send Message", type="primary", use_container_width=True)
                with col_clear:
                    clear_button = st.form_submit_button("üóëÔ∏è Clear Chat", use_container_width=True)
                
                if send_button and user_message.strip():
                    # Add doctor message to history
                    st.session_state.chat_history.append({
                        "role": "doctor",
                        "content": user_message.strip(),
                        "timestamp": datetime.now()
                    })
                    
                    # Get AI response with progress indicator
                    with st.spinner("ü§ñ AI is thinking..."):
                        try:
                            # Check if autogen_agent is initialized
                            if 'autogen_agent' not in st.session_state or st.session_state.autogen_agent is None:
                                # Try to initialize it
                                if st.session_state.get('use_fastapi', False):
                                    st.session_state.autogen_agent = AutoGenMedicalAgent(
                                        st.session_state.rag_system if 'rag_system' in st.session_state else None,
                                        st.session_state.api_client
                                    )
                                else:
                                    if 'rag_system' not in st.session_state:
                                        st.error("‚ùå RAG system not initialized. Please refresh the page.")
                                        st.stop()
                                    st.session_state.autogen_agent = AutoGenMedicalAgent(st.session_state.rag_system, None)
                            
                            start_time = time.time()
                            ai_response = st.session_state.autogen_agent.chat_with_doctor(
                                user_message.strip(), 
                                st.session_state.current_patient
                            )
                            elapsed = time.time() - start_time
                            if elapsed < 2:
                                st.success(f"‚ö° Response generated in {elapsed:.1f}s")
                            
                            # Add AI response to history
                            st.session_state.chat_history.append({
                                "role": "ai",
                                "content": ai_response,
                                "timestamp": datetime.now()
                            })
                        except Exception as e:
                            st.session_state.chat_history.append({
                                "role": "ai",
                                "content": f"‚ùå Error: {str(e)}",
                                "timestamp": datetime.now()
                            })
                    
                    st.rerun()
                
                if clear_button:
                    st.session_state.chat_history = []
                    st.rerun()
            
            # Action buttons with modern styling
            st.markdown("<div style='margin-top: 3rem; margin-bottom: 1.5rem;'></div>", unsafe_allow_html=True)
            st.markdown("---")
            st.markdown("""
            <div class="card-header" style="margin-top: 2rem; margin-bottom: 2rem;">
                <div class="card-header-icon">üöÄ</div>
                <h3 style="margin: 0; color: var(--text-dark);">Quick Actions</h3>
            </div>
            """, unsafe_allow_html=True)
            st.markdown("<div style='margin-bottom: 1.5rem;'></div>", unsafe_allow_html=True)
            col_btn1, col_btn2 = st.columns(2)
            
            with col_btn1:
                # Check if patient is preprocessed (from ChromaDB)
                is_preprocessed = st.session_state.current_patient.get('_is_preprocessed', False) or st.session_state.current_patient.get('_source') == 'chromadb'
                
                if is_preprocessed and st.session_state.discharge_summary:
                    # Show "View Discharge Summary" for preprocessed patients
                    if st.button("üìÑ View Discharge Summary", type="primary", use_container_width=True):
                        # Summary is already loaded, just scroll to it or show it
                        st.session_state.show_summary = True
                        st.rerun()
                else:
                    # Show "Generate Summary" for new patients or preprocessed without summary
                    if st.button("üìù Generate Summary", type="primary", use_container_width=True):
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        try:
                            status_text.text("üìù Formatting patient data...")
                            progress_bar.progress(20)
                            patient_text = st.session_state.rag_system.format_patient_fields(st.session_state.current_patient)
                            
                            status_text.text("ü§ñ Generating summary with AI...")
                            progress_bar.progress(40)
                            start_time = time.time()
                            
                            # Use FastAPI if available
                            if st.session_state.get('use_fastapi', False):
                                template_outline = st.session_state.get("template_outline")
                                summary = st.session_state.api_client.generate_summary(patient_text, template_outline)
                            else:
                                # If a template outline exists, follow it strictly
                                if "template_outline" in st.session_state and st.session_state.template_outline:
                                    summary = st.session_state.rag_system.generate_discharge_summary_with_template(patient_text, st.session_state.template_outline)
                                else:
                                    summary = st.session_state.rag_system.generate_discharge_summary(patient_text)
                            
                            elapsed = time.time() - start_time
                            progress_bar.progress(80)
                            status_text.text("üìÑ Preparing document...")
                            
                            st.session_state.discharge_summary = summary
                            # Build PDF (with template if provided)
                            template_bytes = st.session_state.get("template_pdf_bytes", None)
                            # For template mode, generate a clean PDF using the template's page size but avoid overlaying duplicate headings
                            pdf_bytes = st.session_state.rag_system.generate_pdf_from_text(summary, template_bytes=None if st.session_state.get("template_outline") else template_bytes)
                            st.session_state.discharge_summary_pdf = pdf_bytes
                            
                            progress_bar.progress(100)
                            status_text.text(f"‚úÖ Summary generated in {elapsed:.1f}s")
                            time.sleep(0.5)
                            progress_bar.empty()
                            status_text.empty()
                            
                            st.success(f"‚úÖ Discharge summary generated successfully in {elapsed:.1f}s!")
                            st.rerun()
                        except Exception as e:
                            progress_bar.empty()
                            status_text.empty()
                            st.error(f"‚ùå Error generating summary: {str(e)}")
                            import traceback
                            st.debug(traceback.format_exc())
            
            with col_btn2:
                if st.button("üìä Patient Overview", use_container_width=True):
                    st.session_state.show_patient_overview = True
                    st.rerun()
        
        else:
            st.markdown("""
            <div class="empty-state">
                <div class="icon">üëà</div>
                <h3>Search for a Patient</h3>
                <p>Please search for a patient in the sidebar to start the conversation with the AI assistant.</p>
            </div>
            """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="card-header" style="margin-top: 1rem; margin-bottom: 2.5rem;">
            <div class="card-header-icon">üìã</div>
            <h3 style="margin: 0; color: var(--text-dark);">Generated Content</h3>
        </div>
        """, unsafe_allow_html=True)
        
        # Display discharge summary (editable)
        if st.session_state.discharge_summary:
            st.markdown("""
            <div class="summary-card" style="margin-top: 1.5rem; margin-bottom: 2rem;">
                <div class="card-header" style="margin-top: 0; margin-bottom: 1.5rem;">
                    <div class="card-header-icon">üìÑ</div>
                    <h3 style="margin: 0; color: var(--text-dark);">Discharge Summary (Editable)</h3>
                </div>
            </div>
            """, unsafe_allow_html=True)

            if "editable_summary" not in st.session_state:
                st.session_state.editable_summary = st.session_state.discharge_summary
            # Use current discharge_summary if editable_summary is empty or reset
            elif not st.session_state.editable_summary:
                st.session_state.editable_summary = st.session_state.discharge_summary

            st.session_state.editable_summary = st.text_area(
                 "editable_summary",
                 value=st.session_state.editable_summary,
                 height=500,
                 label_visibility="collapsed"
             )
            
            st.markdown("<div style='margin-top: 2rem; margin-bottom: 1rem;'></div>", unsafe_allow_html=True)

            col_save, col_reset = st.columns([1,1])
            with col_save:
                 if st.button("üíæ Save Edits", use_container_width=True):
                     st.session_state.discharge_summary = st.session_state.editable_summary
                     st.success("Saved your edits.")
            with col_reset:
                 if st.button("‚Ü©Ô∏è Reset to Generated", use_container_width=True):
                    st.session_state.editable_summary = st.session_state.discharge_summary
                    st.rerun() # Rerun to ensure text_area updates

            # --- START: NEW FEEDBACK LOOP UI ---
            st.markdown("<div style='margin-top: 2.5rem; margin-bottom: 1.5rem;'></div>", unsafe_allow_html=True)
            st.markdown("---")
            st.markdown("<div style='margin-top: 1.5rem; margin-bottom: 1rem;'></div>", unsafe_allow_html=True)
            st.markdown("### üß† RAG Feedback Loop")
            
            if st.button("Commit Summary to Knowledgebase", 
                         type="primary", 
                         use_container_width=True, 
                         help="Embed this summary and add it to the RAG system for future 'similar cases' searches."):
                
                if st.session_state.editable_summary and st.session_state.current_patient:
                    with st.spinner("Embedding summary and updating knowledgebase..."):
                        st.session_state.rag_system.add_summary_to_vector_db(
                            st.session_state.current_patient,
                            st.session_state.editable_summary
                        )
                else:
                    st.warning("Please ensure a patient is loaded and a summary is present.")
            # --- END: NEW FEEDBACK LOOP UI ---
            
            st.markdown("---") # Added a separator

            # Plain text download
            st.download_button(
                label="üì• Download as .txt",
                data=st.session_state.editable_summary,
                file_name=f"discharge_summary_{st.session_state.current_patient.get('unit no', 'unknown')}.txt",
                mime="text/plain"
            )

            # DOCX download
            docx_bytes = st.session_state.rag_system.generate_docx_from_text(st.session_state.editable_summary)
            st.download_button(
                label="üìù Download as .docx",
                data=docx_bytes,
                file_name=f"discharge_summary_{st.session_state.current_patient.get('unit no', 'unknown')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # PDF download (optional template mode)
            if "discharge_summary_pdf" in st.session_state and st.session_state.discharge_summary_pdf:
                st.download_button(
                    label="üßæ Download PDF (Template Applied)",
                    data=st.session_state.discharge_summary_pdf,
                    file_name=f"discharge_summary_{st.session_state.current_patient.get('unit no', 'unknown')}.pdf",
                    mime="application/pdf"
                )
    
    # Patient Overview Modal (displayed at top when requested)
    if st.session_state.show_patient_overview and st.session_state.current_patient:
        patient = st.session_state.current_patient
        
        # Helper function to format text with standardized subheadings
        def format_text_with_subheadings(text: str) -> str:
            """Format text by making subheadings smaller and professional"""
            import re
            if not text or text == 'N/A':
                return text
            
            # List of common subheadings to standardize
            subheadings = [
                r'SECONDARY DIAGNOSES?:',
                r'ACUTE ISSUES?:',
                r'CHRONIC ISSUES?:',
                r'PRIMARY DIAGNOSIS:',
                r'DISCHARGE DIAGNOSIS:',
                r'ADMISSION DIAGNOSIS:',
                r'PROCEDURES?:',
                r'MEDICATIONS?:',
                r'ALLERGIES?:',
                r'COMPLICATIONS?:',
                r'FOLLOW-UP:',
                r'DISCHARGE INSTRUCTIONS?:',
            ]
            
            formatted_text = str(text)
            # Replace large subheadings with smaller, standardized ones
            for pattern in subheadings:
                # Match the pattern (case insensitive)
                regex = re.compile(pattern, re.IGNORECASE)
                # Replace with smaller heading
                formatted_text = regex.sub(
                    lambda m: f'<span style="font-size: 0.95rem; font-weight: 600; color: var(--primary-light);">{m.group(0).strip()}</span>',
                    formatted_text
                )
            
            return formatted_text
        
        # Create modal-like container
        st.markdown("""
        <div style="background: var(--bg-card); border: 2px solid var(--primary); border-radius: 24px; padding: 2rem; margin: 1rem 0; box-shadow: var(--shadow-lg);">
        """, unsafe_allow_html=True)
        
        col_title, col_close = st.columns([4, 1])
        with col_title:
            st.markdown("### üìä Patient Overview")
        with col_close:
            # Use form for faster response
            with st.form(key="close_overview_form"):
                if st.form_submit_button("‚úï Close", use_container_width=True):
                    st.session_state.show_patient_overview = False
                    st.rerun()
        
        st.markdown("---")
        
        # Basic Information
        col1, col2 = st.columns(2)
        with col1:
            st.markdown(f"""
            <div style="background: var(--bg-secondary); border-radius: 12px; padding: 1rem; margin: 0.5rem 0;">
                <p style="margin: 0.25rem 0; color: var(--text-primary); font-size: 0.95rem;"><strong>Name:</strong> {patient.get('name', 'Unknown')}</p>
                <p style="margin: 0.25rem 0; color: var(--text-primary); font-size: 0.95rem;"><strong>Unit No:</strong> {patient.get('unit no', 'N/A')}</p>
                <p style="margin: 0.25rem 0; color: var(--text-primary); font-size: 0.95rem;"><strong>Date of Birth:</strong> {patient.get('date of birth', 'N/A')}</p>
            </div>
            """, unsafe_allow_html=True)
        with col2:
            st.markdown(f"""
            <div style="background: var(--bg-secondary); border-radius: 12px; padding: 1rem; margin: 0.5rem 0;">
                <p style="margin: 0.25rem 0; color: var(--text-primary); font-size: 0.95rem;"><strong>Sex:</strong> {patient.get('sex', 'N/A')}</p>
                <p style="margin: 0.25rem 0; color: var(--text-primary); font-size: 0.95rem;"><strong>Service:</strong> {patient.get('service', 'N/A')}</p>
                <p style="margin: 0.25rem 0; color: var(--text-primary); font-size: 0.95rem;"><strong>Admission Date:</strong> {patient.get('admission date', 'N/A')}</p>
            </div>
            """, unsafe_allow_html=True)
        
        # Clinical Information
        st.markdown(f"""
        <div style="background: var(--bg-secondary); border-radius: 12px; padding: 1rem; margin: 0.5rem 0;">
            <p style="margin: 0.25rem 0; color: var(--text-primary); font-size: 0.95rem;"><strong>Chief Complaint:</strong> {patient.get('chief complaint', 'N/A')}</p>
            <p style="margin: 0.25rem 0; color: var(--text-primary); font-size: 0.95rem;"><strong>Attending Physician:</strong> {patient.get('attending', 'N/A')}</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Past Medical History
        pmh = patient.get('past medical history', 'N/A')
        if pmh and pmh != 'N/A' and str(pmh).strip():
            formatted_pmh = format_text_with_subheadings(str(pmh))
            st.markdown(f"""
            <div style="background: var(--bg-secondary); border-radius: 12px; padding: 1rem; margin: 0.5rem 0;">
                <p style="margin: 0 0 0.5rem 0; color: var(--primary-light); font-weight: 600; font-size: 1rem;">üìú Past Medical History</p>
                <div style="color: var(--text-secondary); line-height: 1.6; font-size: 0.9rem;">{formatted_pmh}</div>
            </div>
            """, unsafe_allow_html=True)
        
        # Allergies
        allergies = patient.get('allergies', 'N/A')
        if allergies and allergies != 'N/A':
            formatted_allergies = format_text_with_subheadings(str(allergies))
            st.markdown(f"""
            <div style="background: var(--bg-secondary); border-radius: 12px; padding: 1rem; margin: 0.5rem 0;">
                <p style="margin: 0 0 0.5rem 0; color: var(--primary-light); font-weight: 600; font-size: 1rem;">‚ö†Ô∏è Allergies</p>
                <div style="color: var(--text-secondary); font-size: 0.9rem;">{formatted_allergies}</div>
            </div>
            """, unsafe_allow_html=True)
        
        # Procedures
        procedure = patient.get('major surgical or invasive procedure', 'N/A')
        if procedure and procedure != 'N/A' and str(procedure).strip():
            formatted_procedure = format_text_with_subheadings(str(procedure))
            st.markdown(f"""
            <div style="background: var(--bg-secondary); border-radius: 12px; padding: 1rem; margin: 0.5rem 0;">
                <p style="margin: 0 0 0.5rem 0; color: var(--primary-light); font-weight: 600; font-size: 1rem;">üî¨ Procedures</p>
                <div style="color: var(--text-secondary); line-height: 1.6; font-size: 0.9rem;">{formatted_procedure}</div>
            </div>
            """, unsafe_allow_html=True)
        
        # Discharge Diagnosis
        dx = patient.get('discharge diagnosis', 'N/A')
        if dx and dx != 'N/A' and str(dx).strip():
            formatted_dx = format_text_with_subheadings(str(dx))
            st.markdown(f"""
            <div style="background: var(--bg-secondary); border-radius: 12px; padding: 1rem; margin: 0.5rem 0;">
                <p style="margin: 0 0 0.5rem 0; color: var(--primary-light); font-weight: 600; font-size: 1rem;">ü©∫ Discharge Diagnosis</p>
                <div style="color: var(--text-secondary); line-height: 1.6; font-size: 0.9rem;">{formatted_dx}</div>
            </div>
            """, unsafe_allow_html=True)
        
        # Brief Hospital Course
        course = patient.get('brief hospital course', 'N/A')
        if course and course != 'N/A' and str(course).strip():
            formatted_course = format_text_with_subheadings(str(course))
            st.markdown(f"""
            <div style="background: var(--bg-secondary); border-radius: 12px; padding: 1rem; margin: 0.5rem 0;">
                <p style="margin: 0 0 0.5rem 0; color: var(--primary-light); font-weight: 600; font-size: 1rem;">üìù Brief Hospital Course</p>
                <div style="color: var(--text-secondary); line-height: 1.6; font-size: 0.9rem;">{formatted_course}</div>
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("</div>", unsafe_allow_html=True)
    
    # Footer removed - status is now in header

if __name__ == "__main__":
    main()
